"""
PDF Text Extraction Module
Uses PyMuPDF for native text extraction with PaddleOCR fallback for scanned documents
NOW SUPPORTS: PDF, Word (DOCX/DOC), Text, Markdown, RTF, ODT
"""

import fitz  # PyMuPDF
from paddleocr import PaddleOCR
from PIL import Image
import pdfplumber
import io
import numpy as np
from typing import Tuple, List, Dict, Union, Optional
import os
import chardet
import time
import logging
from app import config
from app.cache import get_cache
from app.large_file_handler import LargePDFHandler, should_use_large_file_handler

logger = logging.getLogger(__name__)

PYMUPDF4LLM_AVAILABLE = False
try:
    import pymupdf4llm
    PYMUPDF4LLM_AVAILABLE = True
    print("[OK] PyMuPDF4LLM available for enhanced extraction")
except ImportError:
    print("[WARNING]  PyMuPDF4LLM not installed - using standard extraction only")

# Import document format handlers
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


class PDFExtractor:
    """Extracts text from multiple document formats (PDF, Word, Text, Markdown, etc.)"""
    
    def __init__(self):
        """Initialize document extractor with OCR engine"""
        self.ocr = None  # Lazy loading for OCR
        self.use_enhanced = config.USE_ENHANCED_EXTRACTION and PYMUPDF4LLM_AVAILABLE

        # Extraction statistics
        self.extraction_stats = {
            'enhanced_success': 0,
            'enhanced_failure': 0,
            'fallback_used': 0,
            'ocr_used': 0
        }

        if self.use_enhanced:
            print(f"[OK] Structured extraction: ENABLED (Strategy: {config.PYMUPDF4LLM_TABLE_STRATEGY})")
        else:
            print("[OK] Standard extraction: ENABLED")
        
    def _initialize_ocr(self):
        """Initialize PaddleOCR (lazy loading to save memory)"""
        if self.ocr is None:
            self.ocr = PaddleOCR(
                use_angle_cls = False,
                lang=config.PADDLEOCR_LANG,
                use_gpu=config.PADDLEOCR_USE_GPU,
                show_log=config.PADDLEOCR_SHOW_LOG
            )
    
    def _get_file_extension(self, file_path: str) -> str:
        """Get file extension in lowercase"""
        return os.path.splitext(file_path)[1].lower().replace('.', '')
    
    def _detect_encoding(self, file_path: str) -> str:
        """Detect file encoding for text files"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)
                result = chardet.detect(raw_data)
                return result['encoding'] if result['encoding'] else 'utf-8'
        except:
            return 'utf-8'
    
    # ============================================================
    # STRUCTURED EXTRACTION (PyMuPDF4LLM)
    # ============================================================

    def extract_table_text_from_pages(self, pdf_path: str, page_indices: List[int]) -> Dict[int, str]:
        """
        Extracts tables from specific pages using pdfplumber and converts them 
        to Markdown format for better LLM comprehension.
        
        Args:
            pdf_path: Path to the PDF file
            page_indices: List of 1-based page numbers to process
            
        Returns:
            Dict mapping page number to extracted Markdown table string
        """
        refined_pages = {}
        
        if not page_indices or not os.path.exists(pdf_path):
            return refined_pages

        try:
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
                
                for page_num in page_indices:
                    # Convert 1-based index to 0-based for pdfplumber
                    zero_based_index = page_num - 1
                    
                    if zero_based_index < 0 or zero_based_index >= total_pages:
                        continue
                        
                    page = pdf.pages[zero_based_index]
                    
                    # Extract tables with permissive settings
                    tables = page.extract_tables({
                        "vertical_strategy": "text", 
                        "horizontal_strategy": "text",
                        "intersection_y_tolerance": 10
                    })
                    
                    if not tables:
                        continue
                        
                    # Convert tables to Markdown format
                    md_tables_text = []
                    for table in tables:
                        # Clean the table: filter None and empty strings
                        clean_table = [[(cell or "").strip().replace("\n", " ") for cell in row] for row in table]
                        
                        # Skip tables that are too small or empty
                        if not clean_table or len(clean_table) < 2:
                            continue

                        # Generate Markdown Table
                        # 1. Header
                        header = "| " + " | ".join(clean_table[0]) + " |"
                        # 2. Separator
                        separator = "| " + " | ".join(["---"] * len(clean_table[0])) + " |"
                        # 3. Body
                        body_rows = []
                        for row in clean_table[1:]:
                            body_rows.append("| " + " | ".join(row) + " |")
                        
                        full_table = f"\n\n**[Extracted Table - Page {page_num}]**\n" + \
                                     f"{header}\n{separator}\n" + \
                                     "\n".join(body_rows) + "\n"
                        
                        md_tables_text.append(full_table)
                    
                    if md_tables_text:
                        refined_pages[page_num] = "\n".join(md_tables_text)
                        
        except Exception as e:
            logger.error(f"Table extraction failed for {pdf_path}: {e}")
            
        return refined_pages

    def _extract_pdf_enhanced(self, pdf_path: str) -> Dict[int, str]:
        """
        Extract PDF with Markdown structure using PyMuPDF4LLM
        Returns: {page_number: markdown_text}
        """
        if not PYMUPDF4LLM_AVAILABLE:
            raise ImportError("PyMuPDF4LLM not available")
        
        print(f"  🔍 Structured extraction (Strategy: {config.PYMUPDF4LLM_TABLE_STRATEGY})")
        start_time = time.time()
        
        try:
           # Convert PDF to Markdown
            md_text = pymupdf4llm.to_markdown(
                pdf_path,
                page_chunks=config.PYMUPDF4LLM_PAGE_CHUNKS,
                write_images=config.PYMUPDF4LLM_WRITE_IMAGES,
                table_strategy=config.PYMUPDF4LLM_TABLE_STRATEGY,
                extract_words=config.PYMUPDF4LLM_EXTRACT_WORDS
            )
            # Parse result based on page_chunks setting
            pages_dict = {}
            if config.PYMUPDF4LLM_PAGE_CHUNKS and isinstance(md_text, list):
                # Result is list of page dictionaries
                for page_data in md_text:
                    if isinstance(page_data, dict):
                        page_num = page_data.get('metadata', {}).get('page', len(pages_dict) + 1)
                        page_text = page_data.get('text', '')
                    else:
                        # Fallback: treat as string
                        page_num = len(pages_dict) + 1
                        page_text = str(page_data)
                    
                    if page_text.strip():
                        pages_dict[page_num] = page_text
            else:
                # Result is single string - split by page markers
                if isinstance(md_text, str):
                    # Try to split by page markers
                    if '\n-----\n' in md_text:
                        # PyMuPDF4LLM page separator
                        page_texts = md_text.split('\n-----\n')
                        for i, page_text in enumerate(page_texts, 1):
                            if page_text.strip():
                                pages_dict[i] = page_text.strip()
                    else:
                        # Single page or no separators
                        pages_dict[1] = md_text
                else:
                    # Fallback
                    pages_dict[1] = str(md_text)
            
            elapsed = time.time() - start_time
            print(f"  [OK] Enhanced extraction: {len(pages_dict)} pages in {elapsed:.2f}s")
            
            if 'enhanced_success' in self.extraction_stats:
                self.extraction_stats['enhanced_success'] += 1

            return pages_dict
            
        except Exception as e:
            print(f"  [WARNING]  Enhanced extraction failed: {e}")
            if 'enhanced_failure' in self.extraction_stats:
                self.extraction_stats['enhanced_failure'] += 1
            
            if config.EXTRACTION_FALLBACK_ENABLED:
                print(f"  ↻ Falling back to standard extraction")
                if 'fallback_used' in self.extraction_stats:
                    self.extraction_stats['fallback_used'] += 1
                return self._extract_pdf_pages(pdf_path)
            else:
                raise

    # ============================================================
    # STANDARD EXTRACTION (PyMuPDF + PaddleOCR)
    # ============================================================

    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from plain text files
        
        Args:
            file_path: Path to text file
            
        Returns:
            Extracted text content
        """
        try:
            encoding = self._detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                text = f.read()
            return text
        except Exception as e:
            print(f"Text extraction error: {e}")
            return ""
    
    def extract_text_from_markdown(self, file_path: str) -> str:
        """
        Extract text from Markdown files
        
        Args:
            file_path: Path to markdown file
            
        Returns:
            Extracted text content
        """
        try:
            encoding = self._detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                md_content = f.read()
            
            # Simple text cleanup (no markdown library needed)
            import re
            
            # Remove headers (# ## ###)
            text = re.sub(r'^#{1,6}\s+', '', md_content, flags=re.MULTILINE)
            
            # Remove bold/italic markers
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # **bold**
            text = re.sub(r'__([^_]+)__', r'\1', text)      # __bold__
            text = re.sub(r'\*([^*]+)\*', r'\1', text)      # *italic*
            text = re.sub(r'_([^_]+)_', r'\1', text)        # _italic_
            
            # Remove links but keep link text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            
            # Remove inline code backticks
            text = re.sub(r'`([^`]+)`', r'\1', text)
            
            # Remove code blocks
            text = re.sub(r'``````', '', text, flags=re.DOTALL)
            
            # Remove image syntax
            text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', '', text)
            
            # Clean up excessive whitespace
            text = re.sub(r'\n\s*\n', '\n\n', text)
            
            return text.strip()
            
        except Exception as e:
            print(f"Markdown extraction error: {e}")
            # Fallback: return raw content
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            except:
                return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from Word DOCX files
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        if DocxDocument is None:
            print("python-docx not installed. Cannot extract from DOCX.")
            return ""
        
        try:
            # Validate file exists and is not empty
            if not os.path.exists(file_path):
                print(f"DOCX file not found: {file_path}")
                return ""
            
            if os.path.getsize(file_path) == 0:
                print(f"DOCX file is empty: {file_path}")
                return ""

            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            result = '\n'.join(text_content)
            if not result.strip():
                print(f"[WARNING]  Warning: No text content found in DOCX: {file_path}")
                return ""
            
            return result
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            # Fallback: Try to read as ZIP and extract XML
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                with zipfile.ZipFile(file_path) as docx_zip:
                    # Read document.xml content
                    xml_content = docx_zip.read('word/document.xml')
                    root = ET.fromstring(xml_content)
                    
                    # Extract text from all text nodes
                    texts = []
                    for elem in root.iter():
                        if elem.text:
                            texts.append(elem.text)
                    
                    return '\n'.join(texts)
            except Exception as fallback_error:
                print(f"DOCX fallback extraction also failed: {fallback_error}")
                return ""
    
    def extract_text_from_rtf(self, file_path: str) -> str:
        """
        Extract text from RTF files (basic extraction)
        
        Args:
            file_path: Path to RTF file
            
        Returns:
            Extracted text content
        """
        try:
            import re
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            # Basic RTF to text conversion (remove RTF control words)
            text = re.sub(r'\\{[^}]*\\}', '', rtf_content)
            text = re.sub(r'\\\\[a-z]+\\d*\\s?', '', text)
            text = re.sub(r'[{}]', '', text)
            return text.strip()
        except Exception as e:
            print(f"RTF extraction error: {e}")
            return ""
    
    def extract_text_native(self, pdf_path: str) -> Tuple[str, bool]:
        """
        Extract text using PyMuPDF native extraction (PDF only)
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Tuple of (extracted_text, is_successful)
        """
        try:
            doc = fitz.open(pdf_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                text_content.append(text)
            
            doc.close()
            full_text = "\n".join(text_content)
            
            # Check if extraction was successful (enough text)
            is_successful = len(full_text.strip()) > config.PDF_TEXT_THRESHOLD
            
            return full_text, is_successful
            
        except Exception as e:
            print(f"Native extraction error: {e}")
            return "", False
    
    def extract_text_ocr(self, pdf_path: str) -> str:
        """
        Extract text using PaddleOCR for scanned documents (PDF only)
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text from OCR
        """
        self._initialize_ocr()
        self.extraction_stats['ocr_used'] += 1

        try:
            doc = fitz.open(pdf_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                img_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(img_data))
                
                # Convert to numpy array
                img_array = np.array(image)
                
                # Perform OCR
                result = self.ocr.ocr(img_array)
                
                # Extract text from OCR results
                page_text = []
                if result and result[0]:
                    for line in result[0]:
                        if len(line) >= 2:
                            text, confidence = line[1]
                            if confidence > config.OCR_CONFIDENCE_THRESHOLD:
                                page_text.append(text)
                
                text_content.append(" ".join(page_text))
            
            doc.close()
            return "\\n".join(text_content)
            
        except Exception as e:
            print(f"OCR extraction error: {e}")
            return ""
    
    def extract_pages(self, file_path: str) -> Dict[int, str]:
        """
        Extract text content split by pages
        
        Optimizations applied:
        - Caching (Opt #4)
        - Chunked processing for large files (Opt #3)
        
        Args:
            file_path: Path to document
        
        Returns:
            Dictionary {page_number: text_content}
        """
        # OPTIMIZATION #4: Check cache first
        cache = get_cache()
        file_hash = cache.get_file_hash(file_path)
        
        cached_pages = cache.get_extracted_text(file_hash)
        if cached_pages:
            logger.info(f"[OK] Cache hit for {os.path.basename(file_path)}")
            return cached_pages
        
        # OPTIMIZATION #3: Use chunked processing for large files
        if should_use_large_file_handler(file_path):
            logger.info(f"Using chunked processing for large file: {os.path.basename(file_path)}")
            handler = LargePDFHandler()
            
            all_pages = {}
            for chunk in handler.extract_pages_chunked(file_path):
                all_pages.update(chunk)
            
            # Cache result
            cache.set_extracted_text(file_hash, all_pages)
            return all_pages
        
        # Normal extraction for regular files
        file_ext = self._get_file_extension(file_path)
        
        if file_ext == 'pdf':
            if self.use_enhanced:
                try:
                    pages = self._extract_pdf_enhanced(file_path)
                except Exception as e:
                    logger.warning(f"Enhanced extraction failed: {e}")
                    if config.EXTRACTION_FALLBACK_ENABLED:
                        pages = self._extract_pdf_pages(file_path)
                    else:
                        raise
            else:
                pages = self._extract_pdf_pages(file_path)
        elif file_ext in ['docx', 'doc']:
            text = self.extract_text_from_docx(file_path)
            pages = self._chunk_text(text) if text else {}
        elif file_ext in ['md', 'markdown']:
            text = self.extract_text_from_markdown(file_path)
            pages = self._chunk_text(text) if text else {}
        elif file_ext == 'txt':
            text = self.extract_text_from_txt(file_path)
            pages = self._chunk_text(text) if text else {}
        else:
            text = self.extract_text(file_path)
            pages = self._chunk_text(text) if text else {}
        
        # Cache result
        cache.set_extracted_text(file_hash, pages)
        
        return pages
    
    def _extract_pdf_pages(self, pdf_path: str) -> Dict[int, str]:
        """Extract PDF text page by page with OCR fallback"""
        pages = {}
        try:
            doc = fitz.open(pdf_path)
            
            for page_num, page in enumerate(doc):
                # 1. Try fast native extraction
                text = page.get_text("text")
                
                # 2. Fallback to OCR if page is essentially empty (scanned)
                # config.PDF_TEXT_THRESHOLD is usually 50 characters
                if len(text.strip()) < config.PDF_TEXT_THRESHOLD:
                    print(f"   ↻ Page {page_num + 1} looks scanned. Running OCR...")
                    
                    # Initialize OCR engine only if needed
                    if self.ocr is None:
                        self._initialize_ocr()
                        
                    try:
                        # Convert page to image for OCR (2x zoom for better quality)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        img_data = pix.tobytes("png")
                        image = Image.open(io.BytesIO(img_data))
                        img_array = np.array(image)
                        
                        # Perform OCR
                        result = self.ocr.ocr(img_array)
                        
                        if result and result[0]:
                            ocr_lines = []
                            for line in result[0]:
                                if len(line) >= 2:
                                    text_str, confidence = line[1]
                                    if confidence > config.OCR_CONFIDENCE_THRESHOLD:
                                        ocr_lines.append(text_str)
                            text = "\n".join(ocr_lines)
                            self.extraction_stats['ocr_used'] += 1
                        
                    except Exception as e:
                        print(f"   [WARNING] OCR failed for page {page_num + 1}: {e}")
                        # Keep whatever little text we found natively
                        pass
                
                pages[page_num + 1] = text
                
            doc.close()
            
        except Exception as e:
            print(f"Error extracting PDF pages: {e}")
            return {1: ""}
            
        return pages
    
    def _chunk_text(self, text: str, chunk_size: int = 3000) -> Dict[int, str]:
        """Helper to split raw text into pseudo-pages"""
        pages = {}
        if not text:
            return pages
            
        chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
        for i, chunk in enumerate(chunks):
            pages[i + 1] = chunk
        return pages

    def extract_text(self, file_path: str) -> str:
        """
        Main extraction method with automatic format detection and fallback
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
        """
        file_ext = self._get_file_extension(file_path)
        filename = os.path.basename(file_path)
        
        # Route to appropriate extractor based on file type
        if file_ext == 'pdf':
            # Try native extraction first
            text, is_successful = self.extract_text_native(file_path)
            
            if is_successful:
                print(f"[OK] Native extraction successful for {filename}")
                return text
            else:
                # Fallback to OCR
                print(f"↻ Using OCR for {filename} (low text content detected)")
                return self.extract_text_ocr(file_path)
        
        elif file_ext in ['docx', 'doc']:
            print(f"[OK] Extracting from Word document: {filename}")
            return self.extract_text_from_docx(file_path)
        
        elif file_ext == 'txt':
            print(f"[OK] Extracting from text file: {filename}")
            return self.extract_text_from_txt(file_path)
        
        elif file_ext in ['md', 'markdown']:
            print(f"[OK] Extracting from Markdown file: {filename}")
            return self.extract_text_from_markdown(file_path)
        
        elif file_ext == 'rtf':
            print(f"[OK] Extracting from RTF file: {filename}")
            return self.extract_text_from_rtf(file_path)
        
        elif file_ext == 'odt':
            print(f"⚠ ODT format detected. Attempting text extraction: {filename}")
            # Try to read as zip file (ODT is XML-based)
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                with zipfile.ZipFile(file_path) as zf:
                    content_xml = zf.read('content.xml')
                    root = ET.fromstring(content_xml)
                    # Extract all text nodes
                    text_nodes = []
                    for elem in root.iter():
                        if elem.text:
                            text_nodes.append(elem.text)
                    return '\\n'.join(text_nodes)
            except Exception as e:
                print(f"ODT extraction error: {e}")
                return ""
        
        else:
            print(f"⚠ Unsupported file format: {file_ext}")
            # Try to read as plain text as fallback
            return self.extract_text_from_txt(file_path)
    
    def extract_multiple(self, file_paths: List[str]) -> dict:
        """
        Extract text from multiple document files
        
        Args:
            file_paths: List of document file paths
            
        Returns:
            Dictionary mapping filename to extracted text
        """
        results = {}
        for file_path in file_paths:
            filename = os.path.basename(file_path)
            text = self.extract_text(file_path)
            results[filename] = text
        return results